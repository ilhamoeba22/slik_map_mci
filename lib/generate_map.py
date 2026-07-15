import sys
import os
import json
import docx
from docx.shared import Pt, RGBColor
import tempfile
import copy

def replace_in_paragraphs(paragraphs, data):
    for p in paragraphs:
        for key, val in data.items():
            placeholder = f"{{{{{key}}}}}"
            if placeholder in p.text:
                for run in p.runs:
                    if placeholder in run.text:
                        run.text = run.text.replace(placeholder, str(val))

def replace_placeholders(doc, data):
    # 1. Replace in standard document paragraphs
    replace_in_paragraphs(doc.paragraphs, data)
    
    # 2. Replace in tables (headers, cells)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_in_paragraphs(cell.paragraphs, data)

def insert_row_before(table, template_row, before_tr):
    # Deepcopy the XML element of the template row
    tr = copy.deepcopy(template_row._tr)
    before_tr.getparent().insertBefore(tr, before_tr)
    return docx.table._Row(tr, table)

def make_cell_bold(cell, bold):
    for p in cell.paragraphs:
        if len(p.runs) > 0:
            for run in p.runs:
                run.bold = bold
        else:
            # If no runs exist, add a run to set bold
            txt = p.text
            p.text = ""
            run = p.add_run(txt)
            run.bold = bold

def style_table_labels_and_values(table):
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            if len(row.cells) == 4:
                if idx in [0, 2]:
                    make_cell_bold(cell, True)
                elif idx in [1, 3]:
                    make_cell_bold(cell, False)

def set_cell_text_preserve_style(cell, text, bold=None):
    if len(cell.paragraphs) > 0:
        p = cell.paragraphs[0]
        if len(p.runs) > 0:
            p.runs[0].text = str(text)
            if bold is not None:
                p.runs[0].bold = bold
            # Clear any other runs in this paragraph
            for r in p.runs[1:]:
                r.text = ""
        else:
            run = p.add_run(str(text))
            if bold is not None:
                run.bold = bold
    else:
        p = cell.add_paragraph()
        run = p.add_run(str(text))
        if bold is not None:
            run.bold = bold

def populate_slik_row(row, no, nama_bank, baki_debet, angsuran, kol, status):
    vals = [str(no), str(nama_bank), str(baki_debet), str(angsuran), str(kol), str(status)]
    for idx, cell in enumerate(row.cells):
        set_cell_text_preserve_style(cell, vals[idx], bold=False)

def generate_map(payload):
    template_path = os.path.join(os.getcwd(), 'public', 'templates', 'MAP_TEMPLATE.docx')
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Template not found at: {template_path}")
        
    doc = docx.Document(template_path)
    
    # Cast numeric values to float/int to prevent type mismatch errors
    gaji = float(payload.get('gaji_num', 1) or 1)
    angsuran_bprs = float(payload.get('angsuran_num', 0) or 0)
    
    # Calculate angsuran bank lain from active slik records
    total_angsuran_lain = 0
    for rec in payload.get('slik_records_raw', []):
        status_lower = rec.get('kondisi_ket', '').lower()
        if 'lunas' not in status_lower and 'dialihkan' not in status_lower:
            total_angsuran_lain += float(rec.get('angsuran', 0) or 0)
            
    iir_val = round(((angsuran_bprs + total_angsuran_lain) / gaji) * 100)
    payload['IIR'] = f"{iir_val} % dari 95 %"
    
    # Placeholders map (primarily for paragraph level replacements)
    placeholders_data = {
        'NO_REGISTRASI': payload.get('NO_REGISTRASI', ''),
        'NAMA': payload.get('NAMA', ''),
        'NIK': payload.get('NIK', ''),
        'NOPEN': payload.get('NOPEN', ''),
        'GAJI': payload.get('GAJI', ''),
        'NAMA_PENSIUN': payload.get('NAMA_PENSIUN', ''),
        'ALAMAT': payload.get('ALAMAT', ''),
        'NO_HP': payload.get('NO_HP', ''),
        'SK_TEMPAT': payload.get('SK_TEMPAT', 'YOGYAKARTA'),
        'KANTOR_ASAL': payload.get('KANTOR_ASAL', ''),
        'KANTOR_TUJUAN': payload.get('KANTOR_TUJUAN', ''),
        'SK_TANGGAL': payload.get('SK_TANGGAL', ''),
        'PENSIUNAN_DARI': payload.get('PENSIUNAN_DARI', ''),
        'NO_SK': payload.get('NO_SK', ''),
        'TGL_LAHIR': payload.get('TGL_LAHIR', ''),
        'USIA_SEKARANG': payload.get('USIA_SEKARANG', ''),
        'USIA_LUNAS': payload.get('USIA_LUNAS', ''),
        'PLAFON': payload.get('PLAFON', ''),
        'TENOR': payload.get('TENOR', ''),
        'MARGIN': payload.get('MARGIN', ''),
        'ANGSURAN': payload.get('ANGSURAN', ''),
        'TUJUAN': payload.get('TUJUAN', ''),
        'IIR': payload.get('IIR', '0 % dari 95 %')
    }
    
    # 1. Do standard paragraph placeholders replacement
    replace_placeholders(doc, placeholders_data)

    # 1.3 Update Table 0 (INFORMASI PERMOHONAN PEMBIAYAAN) cells dynamically
    table_0 = doc.tables[0]
    if len(table_0.rows) > 8:
        # Row 0
        set_cell_text_preserve_style(table_0.rows[0].cells[1], payload.get('NO_REGISTRASI', ''))
        set_cell_text_preserve_style(table_0.rows[0].cells[3], payload.get('NO_CIF', '-'))
        # Row 1
        set_cell_text_preserve_style(table_0.rows[1].cells[1], payload.get('JENIS_NASABAH', 'Perorangan'))
        set_cell_text_preserve_style(table_0.rows[1].cells[3], payload.get('NO_REK_PEMBIAYAAN', '-'))
        # Row 2
        set_cell_text_preserve_style(table_0.rows[2].cells[1], payload.get('STATUS_NASABAH', 'Nasabah Baru'))
        set_cell_text_preserve_style(table_0.rows[2].cells[3], payload.get('NO_REK_SIMPANAN', '-'))
        # Row 3
        set_cell_text_preserve_style(table_0.rows[3].cells[1], payload.get('PRODUK_PEMBIAYAAN', 'Multiguna Pensiunan'))
        set_cell_text_preserve_style(table_0.rows[3].cells[3], payload.get('STATUS_PEMBIAYAAN', 'SK Ditangan'))
        # Row 4
        set_cell_text_preserve_style(table_0.rows[4].cells[1], payload.get('MITRA_VENDOR', 'Grahadi'))
        set_cell_text_preserve_style(table_0.rows[4].cells[3], payload.get('ASURANSI', 'Reliance (Asuransi Jiwa)'))
        # Row 5
        set_cell_text_preserve_style(table_0.rows[5].cells[1], payload.get('PKS_NO', '-'))
        set_cell_text_preserve_style(table_0.rows[5].cells[3], payload.get('JENIS_PEMBAYARAN', 'Angsuran Bulanan'))
        # Row 6
        set_cell_text_preserve_style(table_0.rows[6].cells[1], payload.get('JENIS_PEMBIAYAAN', 'Multiguna'))
        set_cell_text_preserve_style(table_0.rows[6].cells[3], payload.get('SKEMA_PEMBAYARAN', 'Flat'))
        # Row 7
        set_cell_text_preserve_style(table_0.rows[7].cells[1], payload.get('JENIS_AKAD', 'Murabahah'))
        set_cell_text_preserve_style(table_0.rows[7].cells[3], payload.get('JENIS_PENGGUNAAN', 'Konsumtif'))
        # Row 8
        set_cell_text_preserve_style(table_0.rows[8].cells[1], payload.get('JENIS_PENGIKATAN', 'Fidusia'))
        set_cell_text_preserve_style(table_0.rows[8].cells[3], payload.get('TUJUAN', ''))

    # 1.5 Update Table 1 (INFORMASI DATA NASABAH) cells dynamically
    table_1 = doc.tables[1]
    if len(table_1.rows) > 13:
        # Row 0
        set_cell_text_preserve_style(table_1.rows[0].cells[1], payload.get('NAMA', ''))
        set_cell_text_preserve_style(table_1.rows[0].cells[3], payload.get('NAMA_PASANGAN', '-'))
        # Row 1
        set_cell_text_preserve_style(table_1.rows[1].cells[1], payload.get('NIK', ''))
        set_cell_text_preserve_style(table_1.rows[1].cells[3], payload.get('NIK_PASANGAN', '-'))
        # Row 2
        set_cell_text_preserve_style(table_1.rows[2].cells[1], payload.get('TGL_LAHIR', ''))
        set_cell_text_preserve_style(table_1.rows[2].cells[3], payload.get('TGL_LAHIR_PASANGAN', '-'))
        # Row 3
        set_cell_text_preserve_style(table_1.rows[3].cells[1], payload.get('USIA_SEKARANG', ''))
        set_cell_text_preserve_style(table_1.rows[3].cells[3], payload.get('JUMLAH_TANGGUNGAN', '0'))
        # Row 4
        set_cell_text_preserve_style(table_1.rows[4].cells[1], payload.get('USIA_LUNAS', ''))
        set_cell_text_preserve_style(table_1.rows[4].cells[3], "Rp. " + payload.get('PENGHASILAN_KOTOR_TAHUN', '0'))
        # Row 5
        set_cell_text_preserve_style(table_1.rows[5].cells[1], payload.get('ALAMAT', ''))
        set_cell_text_preserve_style(table_1.rows[5].cells[3], payload.get('GAJI', ''))
        # Row 6
        set_cell_text_preserve_style(table_1.rows[6].cells[1], payload.get('JENIS_KELAMIN', '-'))
        set_cell_text_preserve_style(table_1.rows[6].cells[3], payload.get('PERJANJIAN_PISAH_HARTA', 'TIDAK'))
        # Row 7
        set_cell_text_preserve_style(table_1.rows[7].cells[1], payload.get('STATUS_PERKAWINAN', ''))
        set_cell_text_preserve_style(table_1.rows[7].cells[3], payload.get('SUMBER_PENGHASILAN', 'GAJI PENSIUN'))
        # Row 8
        set_cell_text_preserve_style(table_1.rows[8].cells[1], payload.get('NO_HP', ''))
        set_cell_text_preserve_style(table_1.rows[8].cells[3], payload.get('NAMA_PENSIUN', ''))
        # Row 9
        set_cell_text_preserve_style(table_1.rows[9].cells[1], 'NASABAH BPRS HIK MCI')
        set_cell_text_preserve_style(table_1.rows[9].cells[3], payload.get('NAMA_PENERIMA_PENSIUN', ''))
        # Row 10
        set_cell_text_preserve_style(table_1.rows[10].cells[1], payload.get('SK_TEMPAT', 'YOGYAKARTA'))
        set_cell_text_preserve_style(table_1.rows[10].cells[3], payload.get('NOPEN', ''))
        # Row 11
        set_cell_text_preserve_style(table_1.rows[11].cells[1], payload.get('KANTOR_ASAL', ''))
        set_cell_text_preserve_style(table_1.rows[11].cells[3], payload.get('NO_SK', ''))
        # Row 12
        set_cell_text_preserve_style(table_1.rows[12].cells[1], payload.get('KANTOR_TUJUAN', ''))
        set_cell_text_preserve_style(table_1.rows[12].cells[3], payload.get('SK_TANGGAL', ''))
        # Row 13
        set_cell_text_preserve_style(table_1.rows[13].cells[1], payload.get('PENSIUNAN_DARI', ''))
        set_cell_text_preserve_style(table_1.rows[13].cells[3], payload.get('STATUS_SK', 'ASLI'))

    # 1.6 Update Table 3 (ANALISA PERHITUNGAN PEMBIAYAAN) cells dynamically
    table_3 = doc.tables[3]
    if len(table_3.rows) > 4:
        # Row 0
        set_cell_text_preserve_style(table_3.rows[0].cells[1], payload.get('GAJI', ''))
        set_cell_text_preserve_style(table_3.rows[0].cells[3], payload.get('BIAYA_ADMINISTRASI', '0'))
        # Row 1
        set_cell_text_preserve_style(table_3.rows[1].cells[1], payload.get('PENSIUNAN_DARI', ''))
        set_cell_text_preserve_style(table_3.rows[1].cells[3], payload.get('ANGSURAN', ''))
        # Row 2
        set_cell_text_preserve_style(table_3.rows[2].cells[1], payload.get('PLAFON', ''))
        set_cell_text_preserve_style(table_3.rows[2].cells[3], payload.get('MARGIN', ''))
        # Row 3
        set_cell_text_preserve_style(table_3.rows[3].cells[1], payload.get('TENOR', ''))
        set_cell_text_preserve_style(table_3.rows[3].cells[3], payload.get('IIR', ''))
        # Row 4
        set_cell_text_preserve_style(table_3.rows[4].cells[1], payload.get('HARGA_JUAL', '0'))
        set_cell_text_preserve_style(table_3.rows[4].cells[3], payload.get('MARGIN_NETT', '0'))

    # Style Table 0, 1, and 3 labels as Bold and values as Not Bold
    style_table_labels_and_values(table_0)
    style_table_labels_and_values(table_1)
    style_table_labels_and_values(table_3)

    # 2. Replace signature block for AO (BELINDA P.J) in paragraphs
    ao_name = payload.get('AO_NAME', 'BELINDA P.J')
    for p in doc.paragraphs:
        if 'BELINDA P.J' in p.text:
            for run in p.runs:
                if 'BELINDA P.J' in run.text:
                    run.text = run.text.replace('BELINDA P.J', ao_name)

    # 3. Replace signature block for Kadiv Bisnis (Faradays Muhammad) in Table 5
    kadiv_name = payload.get('KADIV_BISNIS', 'Faradays Muhammad')
    table_5 = doc.tables[5]
    if len(table_5.rows) > 2 and len(table_5.rows[2].cells) > 0:
        cell = table_5.rows[2].cells[0]
        # Replace either Faradays Muhammad or the placeholder [Nama Kadiv] or [Nama Kadiv] in brackets
        for name_to_replace in ['Faradays Muhammad', '[Nama Kadiv]']:
            if name_to_replace in cell.text:
                for p in cell.paragraphs:
                    for run in p.runs:
                        if name_to_replace in run.text:
                            run.text = run.text.replace(name_to_replace, kadiv_name)

    # 4. Populate Catatan and Keputusan Approval in Table 5
    catatan_text = payload.get('CATATAN_APPROVAL', '')
    keputusan_text = payload.get('KEPUTUSAN_APPROVAL', '')
    if len(table_5.rows) > 3 and len(table_5.rows[0].cells) > 1:
        set_cell_text_preserve_style(table_5.rows[2].cells[1], catatan_text)
        set_cell_text_preserve_style(table_5.rows[3].cells[1], keputusan_text)

    # 5. Modify SLIK table (Table 2)
    slik_table = doc.tables[2]
    
    # Rename 6th column to "Keterangan"
    if len(slik_table.rows) > 0 and len(slik_table.rows[0].cells) > 5:
        set_cell_text_preserve_style(slik_table.rows[0].cells[5], "Keterangan")

    # Find the KETERANGAN header row index in Table 2 dynamically
    keterangan_row_idx = None
    for idx, r in enumerate(slik_table.rows):
        if len(r.cells) > 0 and 'KETERANGAN' in r.cells[0].text:
            keterangan_row_idx = idx
            break
            
    if keterangan_row_idx is None:
        # Fallback if the template structure is different
        keterangan_row_idx = len(slik_table.rows) - 2

    # Keep row 1 as a template row to preserve the exact styles
    template_row = slik_table.rows[1]

    # Remove all data rows between row 1 and the KETERANGAN row
    # Note: initially, in template, rows 1 to 5 are empty data rows.
    # We want to clear them and only keep row 1.
    while keterangan_row_idx > 2:
        slik_table._tbl.remove(slik_table.rows[2]._tr)
        keterangan_row_idx -= 1
        
    records = payload.get('slik_records', [])
    if len(records) == 0:
        # Populate the template row with clean/empty SLIK info
        populate_slik_row(template_row, "1.", "TIDAK ADA FASILITAS AKTIF (SLIK BERSIH)", "0", "0", "Lancar", "Fasilitas Lunas / Bersih")
    else:
        # Populate the first record into the template row
        first_rec = records[0]
        populate_slik_row(
            template_row,
            "1.",
            first_rec['nama_bank'],
            first_rec['baki_debet'],
            first_rec['angsuran'],
            first_rec['kolektibilitas'],
            first_rec['status']
        )
        
        # Insert and populate the rest before the KETERANGAN row
        # The KETERANGAN row is now at index 2
        for idx, rec in enumerate(records[1:]):
            keterangan_row_element = slik_table.rows[2]._tr
            new_row = insert_row_before(slik_table, template_row, keterangan_row_element)
            populate_slik_row(
                new_row,
                f"{idx + 2}.",
                rec['nama_bank'],
                rec['baki_debet'],
                rec['angsuran'],
                rec['kolektibilitas'],
                rec['status']
            )

    # 6. Modify Catatan Analisa table (Table 4)
    catatan_analisa_table = doc.tables[4]
    catatan_analisa_text = payload.get('CATATAN_ANALISA', '').strip()
    if catatan_analisa_text:
        if len(catatan_analisa_table.rows) > 1 and len(catatan_analisa_table.rows[1].cells) > 0:
            set_cell_text_preserve_style(catatan_analisa_table.rows[1].cells[0], catatan_analisa_text)

    # Remove trailing empty or whitespace-only paragraphs to prevent trailing blank pages
    for p in list(reversed(doc.paragraphs)):
        if not p.text.strip():
            p_element = p._element
            p_element.getparent().remove(p_element)
        else:
            break

    # Save to a temporary docx file
    fd, temp_file_path = tempfile.mkstemp(suffix='.docx')
    os.close(fd)
    
    doc.save(temp_file_path)
    return temp_file_path

if __name__ == '__main__':
    try:
        input_data = sys.stdin.read()
        payload = json.loads(input_data)
        
        temp_path = generate_map(payload)
        print(temp_path)
        sys.exit(0)
    except Exception as e:
        print(f"ERROR: {str(e)}", file=sys.stderr)
        sys.exit(1)
