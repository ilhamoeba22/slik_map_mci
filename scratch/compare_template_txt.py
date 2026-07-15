import docx
import os

template_path = r"D:\islamic-banking\Dokumen Grahadi\Template MAP.docx"
if not os.path.exists(template_path):
    print("Template MAP.docx not found!")
    sys.exit(1)

doc = docx.Document(template_path)
print(f"Loaded 'Template MAP.docx' successfully!")
print(f"Total tables: {len(doc.tables)}")

for idx, table in enumerate(doc.tables):
    print(f"\n=== TABLE {idx} ({len(table.rows)} rows, {len(table.columns)} columns) ===")
    for r_idx, row in enumerate(table.rows):
        row_text = []
        for c_idx, cell in enumerate(row.cells):
            # Avoid duplicate printed text for merged cells
            row_text.append(cell.text.strip().replace('\n', ' '))
        print(f"Row {r_idx}:", row_text)
